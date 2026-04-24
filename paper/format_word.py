"""
Format Word document according to Vietnamese academic standards:
- Font: Times New Roman, 13pt body, 14pt headings
- Line spacing: 1.5
- Margins: Left 3cm, Right 2cm, Top 2cm, Bottom 2cm
- First line indent: 1.27cm
- Table: bordered, 11pt, centered
- Caption: italic, 12pt
- Page numbers: bottom center
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document('palm_vein_docs_v3.docx')

# ============================================================
# 1. PAGE SETUP — margins and page numbers
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # Page numbers bottom center
    footer = section.footer
    footer.is_linked_to_previous = False
    if not footer.paragraphs:
        fp = footer.add_paragraph()
    else:
        fp = footer.paragraphs[0]
    fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_xml = (
        '<w:fldSimple w:instr=" PAGE " xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        '<w:sz w:val="26"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
    )
    run._element.getparent().append(parse_xml(fld_xml))

# ============================================================
# 2. STYLE DEFINITIONS
# ============================================================

def set_font(run_or_font, name='Times New Roman', size=None, bold=None, italic=None, color=None):
    """Set font properties on a run or font object."""
    font = run_or_font if hasattr(run_or_font, 'size') and not hasattr(run_or_font, 'text') else run_or_font.font
    font.name = name
    rpr = font.element if hasattr(font, 'element') else font._element
    # Ensure East Asian font is also Times New Roman
    if hasattr(rpr, 'rPr') and rpr.rPr is not None:
        rpr.rPr.set(qn('w:eastAsia'), name)
    if size:
        font.size = size
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color:
        font.color.rgb = color

def set_paragraph_format(para, alignment=None, first_indent=None,
                         space_before=None, space_after=None,
                         line_spacing=None, keep_next=None):
    """Set paragraph formatting."""
    pf = para.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if keep_next is not None:
        pf.keep_with_next = keep_next

def format_all_runs(para, name='Times New Roman', size=None, bold=None, italic=None):
    """Apply font to all runs in a paragraph."""
    for run in para.runs:
        run.font.name = name
        rPr = run._element.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), name)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}" w:cs="{name}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:ascii'), name)
            rFonts.set(qn('w:hAnsi'), name)
            rFonts.set(qn('w:eastAsia'), name)
            rFonts.set(qn('w:cs'), name)
        if size:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            # Only set italic if not already explicitly set differently
            if run.font.italic is None or italic:
                run.font.italic = italic

# ============================================================
# 3. MODIFY BUILT-IN STYLES
# ============================================================

# Normal style
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(13)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
style_normal.paragraph_format.space_before = Pt(0)
style_normal.paragraph_format.space_after = Pt(6)
rPr = style_normal.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman" w:cs="Times New Roman"/>')
existing = rPr.find(qn('w:rFonts'))
if existing is not None:
    rPr.remove(existing)
rPr.insert(0, rFonts)

# Heading 1: 14pt, bold, centered
for style_name in ['Heading 1']:
    if style_name in doc.styles:
        s = doc.styles[style_name]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(14)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        s.paragraph_format.first_line_indent = Cm(0)

# Heading 2: 13pt, bold
for style_name in ['Heading 2']:
    if style_name in doc.styles:
        s = doc.styles[style_name]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(13)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        s.paragraph_format.first_line_indent = Cm(0)

# Heading 3: 13pt, bold italic
for style_name in ['Heading 3']:
    if style_name in doc.styles:
        s = doc.styles[style_name]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(13)
        s.font.bold = True
        s.font.italic = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        s.paragraph_format.first_line_indent = Cm(0)

# ============================================================
# 4. FORMAT EACH PARAGRAPH
# ============================================================

for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name

    if style_name == 'Title':
        # Title: 16pt, bold, centered
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(24), space_after=Pt(12),
                           line_spacing=1.5, first_indent=Cm(0))
        format_all_runs(para, size=Pt(16), bold=True)

    elif style_name == 'Author':
        # Author: 13pt, centered
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(6), space_after=Pt(12),
                           line_spacing=1.5, first_indent=Cm(0))
        format_all_runs(para, size=Pt(13))

    elif style_name == 'Abstract Title':
        # Abstract title: 13pt bold, centered
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(12), space_after=Pt(6),
                           line_spacing=1.5, first_indent=Cm(0))
        format_all_runs(para, size=Pt(13), bold=True)

    elif style_name == 'Abstract':
        # Abstract body: 12pt, italic, justified, indented both sides
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_before=Pt(0), space_after=Pt(6),
                           line_spacing=1.5, first_indent=Cm(1.27))
        format_all_runs(para, size=Pt(12), italic=True)

    elif style_name.startswith('Heading'):
        # Headings: already set via styles, ensure runs match
        level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
        sz = Pt(14) if level == 1 else Pt(13)
        format_all_runs(para, size=sz, bold=True)
        if level >= 3:
            for run in para.runs:
                run.font.italic = True

    elif style_name in ('Table Caption', 'Caption'):
        # Table/Figure caption: 12pt, bold label, centered
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(6), space_after=Pt(3),
                           line_spacing=1.5, first_indent=Cm(0))
        format_all_runs(para, size=Pt(12))
        # Bold the "Bảng X:" or "Hình X:" prefix
        for run in para.runs:
            text = run.text
            if text.startswith(('Bảng', 'Hình', 'Table', 'Figure')):
                run.font.bold = True

    elif style_name == 'Image Caption':
        # Image/Figure caption: 12pt, italic, centered
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(3), space_after=Pt(6),
                           line_spacing=1.5, first_indent=Cm(0))
        format_all_runs(para, size=Pt(12), italic=True)

    elif 'First Paragraph' in style_name:
        # First paragraph after heading: justified, first-line indent
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_before=Pt(0), space_after=Pt(6),
                           line_spacing=1.5, first_indent=Cm(1.27))
        format_all_runs(para, size=Pt(13))

    elif style_name in ('Body Text', 'Normal'):
        # Body text: 13pt, justified, first-line indent
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_before=Pt(0), space_after=Pt(6),
                           line_spacing=1.5, first_indent=Cm(1.27))
        format_all_runs(para, size=Pt(13))

    else:
        # Any other style: apply base formatting
        if para.text.strip():
            set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               line_spacing=1.5)
            format_all_runs(para, size=Pt(13))

# ============================================================
# 5. FORMAT TABLES
# ============================================================

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     size='4', color='000000'):
    """Set borders on a table cell."""
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        tc.insert(0, tcPr)

    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(borders)

    for border_name, enabled in [('top', top), ('bottom', bottom),
                                  ('left', left), ('right', right)]:
        border = borders.find(qn(f'w:{border_name}'))
        if border is not None:
            borders.remove(border)
        if enabled:
            border = parse_xml(
                f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="{size}" '
                f'w:space="0" w:color="{color}"/>'
            )
            borders.append(border)

def set_table_borders(table):
    """Set borders on entire table using tblBorders."""
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remove existing borders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

for table in doc.tables:
    # Center the table
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table borders
    set_table_borders(table)

    # Auto-fit table width
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))

    # Set table width to 100%
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')

    # Format cells
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                # Cell text: 11pt, centered vertically and horizontally
                set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                   space_before=Pt(2), space_after=Pt(2),
                                   line_spacing=1.0, first_indent=Cm(0))
                format_all_runs(para, size=Pt(11))

                # Bold header row (first row)
                if row_idx == 0:
                    for run in para.runs:
                        run.font.bold = True

            # Vertical alignment: center
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                tc.insert(0, tcPr)
            vAlign = tcPr.find(qn('w:vAlign'))
            if vAlign is None:
                vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                tcPr.append(vAlign)
            else:
                vAlign.set(qn('w:val'), 'center')

# ============================================================
# 6. FORMAT BIBLIOGRAPHY
# ============================================================

in_bibliography = False
for para in doc.paragraphs:
    text = para.text.strip()
    if 'Tài liệu tham khảo' in text or text == 'References':
        in_bibliography = True
        continue
    if in_bibliography and text:
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_before=Pt(0), space_after=Pt(3),
                           line_spacing=1.0, first_indent=Cm(-1.27))
        para.paragraph_format.left_indent = Cm(1.27)
        format_all_runs(para, size=Pt(11))

# ============================================================
# 7. SAVE
# ============================================================

output = 'palm_vein_docs_final.docx'
doc.save(output)
print(f'Saved formatted document: {output}')

# Quick verification
doc2 = Document(output)
print(f'Tables: {len(doc2.tables)}')
print(f'Paragraphs: {len(doc2.paragraphs)}')
img_count = sum(1 for r in doc2.part.rels.values() if 'image' in r.reltype)
print(f'Images: {img_count}')
s = doc2.sections[0]
print(f'Margins: L={s.left_margin/914400:.1f}in R={s.right_margin/914400:.1f}in T={s.top_margin/914400:.1f}in B={s.bottom_margin/914400:.1f}in')
print(f'Normal font: {doc2.styles["Normal"].font.name}, {doc2.styles["Normal"].font.size}')
